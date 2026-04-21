"""Resume generation (.docx + .pdf) and fabrication checking.

Three templates, all strictly ATS-safe (single column, standard headings,
no tables/icons/graphics):
  - clean_modern: Calibri (Carlito) 11pt, contemporary (default)
  - harvard: Garamond (EB Garamond) 11pt, centered headings, formal
  - mirror_user: matches the style of the user's uploaded resume

.docx output uses python-docx. .pdf output uses reportlab directly — no
docx→pdf conversion step — so PDFs render identically on every platform
without requiring MS Word, LibreOffice, or Pango/cairo system libs. Bundled
OFL-licensed fonts live in `assets/fonts/` (EB Garamond, Carlito).

Public entry points dispatch on output file extension:
  build_clean_modern(profile, tailored, "resume.docx")  → python-docx
  build_clean_modern(profile, tailored, "resume.pdf")   → reportlab
  build_resume("harvard", profile, tailored, "resume.pdf")

Plus `fabrication_check` which walks the tailored content and verifies every
claim can be traced back to profile.evidence or profile.experience. Stage 6
calls this before presenting any resume.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor


# ---------------------------------------------------------------------------
# Tailored content schema
# ---------------------------------------------------------------------------
#
# The LLM (Stage 6 orchestrator) produces tailored content in this shape
# and passes it to the template builders:
#
# {
#     "name": str,
#     "contact_line": str,                # e.g. "Washington, DC | jane@x.com | (555) 555-5555 | linkedin.com/in/..."
#     "summary": str,                     # 2-4 sentence role-tailored
#     "skills": {
#         "core_competencies": [str, ...],
#         "tools_platforms": [str, ...],
#         "additional": [str, ...],
#     },
#     "experience": [
#         {
#             "title": str, "company": str, "location": str,
#             "start": str, "end": str, "current": bool,
#             "bullets": [str, ...],      # already ordered/elevated/reworded
#         },
#         ...
#     ],
#     "education": [
#         {"degree": str, "major": str, "school": str, "year": int|str, "honors": str|None},
#         ...
#     ],
#     "certifications": [
#         {"name": str, "issuer": str|None, "year": int|str|None},
#         ...
#     ],
# }


# ===========================================================================
# Shared helpers
# ===========================================================================

def _date_range(exp: dict) -> str:
    start = exp.get("start", "") or ""
    end = "Present" if exp.get("current") else (exp.get("end", "") or "")
    if start and end:
        return f"{start} – {end}"
    return start or end or ""


# ===========================================================================
# DOCX BUILDERS (python-docx)
# ===========================================================================

def _set_margins(doc, top=0.7, bottom=0.7, left=0.75, right=0.75):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _add_run(paragraph, text, bold=False, size=11, font_name=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _docx_section_heading(doc, text, font_name="Calibri", size=13,
                          centered=False, underline=False):
    p = doc.add_paragraph()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = font_name
    if underline:
        run.underline = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _build_clean_modern_docx(profile: dict, tailored: dict, out_path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 0.6, 0.6, 0.7, 0.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, tailored["name"], bold=True, size=18, font_name="Calibri")
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    _add_run(p, tailored.get("contact_line", ""), size=11, font_name="Calibri")
    p.paragraph_format.space_after = Pt(10)

    if tailored.get("summary"):
        _docx_section_heading(doc, "SUMMARY", font_name="Calibri", size=13)
        p = doc.add_paragraph()
        _add_run(p, tailored["summary"], size=11, font_name="Calibri")
        p.paragraph_format.space_after = Pt(8)

    if tailored.get("skills"):
        _docx_section_heading(doc, "SKILLS", font_name="Calibri", size=13)
        for label, items in [
            ("Core Competencies", tailored["skills"].get("core_competencies", [])),
            ("Tools & Platforms", tailored["skills"].get("tools_platforms", [])),
            ("Additional", tailored["skills"].get("additional", [])),
        ]:
            if not items:
                continue
            p = doc.add_paragraph()
            _add_run(p, f"{label}: ", bold=True, size=11, font_name="Calibri")
            _add_run(p, ", ".join(items), size=11, font_name="Calibri")
            p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if tailored.get("experience"):
        _docx_section_heading(doc, "EXPERIENCE", font_name="Calibri", size=13)
        for exp in tailored["experience"]:
            p = doc.add_paragraph()
            _add_run(p, exp.get("title", ""), bold=True, size=11, font_name="Calibri")
            _add_run(p, f"  |  {exp.get('company', '')}", size=11, font_name="Calibri")
            if exp.get("location"):
                _add_run(p, f"  |  {exp['location']}", size=11, font_name="Calibri")
            dr = _date_range(exp)
            if dr:
                _add_run(p, f"  |  {dr}", size=11, font_name="Calibri")
            p.paragraph_format.space_after = Pt(2)

            for bullet in exp.get("bullets", []) or []:
                bp = doc.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=11, font_name="Calibri")
                bp.paragraph_format.space_after = Pt(1)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if tailored.get("education"):
        _docx_section_heading(doc, "EDUCATION", font_name="Calibri", size=13)
        for ed in tailored["education"]:
            p = doc.add_paragraph()
            parts = [ed.get("degree", "")]
            if ed.get("major"):
                parts.append(ed["major"])
            left = ", ".join([x for x in parts if x])
            _add_run(p, left, bold=True, size=11, font_name="Calibri")
            _add_run(p, f"  |  {ed.get('school', '')}", size=11, font_name="Calibri")
            if ed.get("year"):
                _add_run(p, f"  |  {ed['year']}", size=11, font_name="Calibri")
            if ed.get("honors"):
                _add_run(p, f"  |  {ed['honors']}", size=11, font_name="Calibri")
            p.paragraph_format.space_after = Pt(2)

    if tailored.get("certifications"):
        _docx_section_heading(doc, "CERTIFICATIONS", font_name="Calibri", size=13)
        for c in tailored["certifications"]:
            p = doc.add_paragraph()
            _add_run(p, c.get("name", ""), bold=True, size=11, font_name="Calibri")
            tail = []
            if c.get("issuer"):
                tail.append(c["issuer"])
            if c.get("year"):
                tail.append(str(c["year"]))
            if tail:
                _add_run(p, "  |  " + "  |  ".join(tail), size=11, font_name="Calibri")

    doc.save(str(out_path))
    return out_path


def _build_harvard_docx(profile: dict, tailored: dict, out_path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 0.75, 0.75, 0.9, 0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, tailored["name"], bold=True, size=16, font_name="Garamond")
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, tailored.get("contact_line", ""), size=11, font_name="Garamond")
    p.paragraph_format.space_after = Pt(10)

    if tailored.get("summary"):
        _docx_section_heading(doc, "SUMMARY", font_name="Garamond", size=12,
                              centered=True, underline=True)
        p = doc.add_paragraph()
        _add_run(p, tailored["summary"], size=11, font_name="Garamond")
        p.paragraph_format.space_after = Pt(6)

    if tailored.get("experience"):
        _docx_section_heading(doc, "EXPERIENCE", font_name="Garamond", size=12,
                              centered=True, underline=True)
        for exp in tailored["experience"]:
            p = doc.add_paragraph()
            _add_run(p, exp.get("company", ""), bold=True, size=11, font_name="Garamond")
            dr = _date_range(exp)
            if dr:
                _add_run(p, "\t" + dr, size=11, font_name="Garamond")
            p.paragraph_format.space_after = Pt(0)

            p = doc.add_paragraph()
            _add_run(p, exp.get("title", ""), size=11, font_name="Garamond")
            if exp.get("location"):
                _add_run(p, "\t" + exp["location"], size=11, font_name="Garamond")
            p.paragraph_format.space_after = Pt(2)

            for bullet in exp.get("bullets", []) or []:
                bp = doc.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=11, font_name="Garamond")
                bp.paragraph_format.space_after = Pt(1)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if tailored.get("education"):
        _docx_section_heading(doc, "EDUCATION", font_name="Garamond", size=12,
                              centered=True, underline=True)
        for ed in tailored["education"]:
            p = doc.add_paragraph()
            _add_run(p, ed.get("school", ""), bold=True, size=11, font_name="Garamond")
            if ed.get("year"):
                _add_run(p, "\t" + str(ed["year"]), size=11, font_name="Garamond")
            p.paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph()
            deg = ", ".join([ed.get("degree", ""), ed.get("major", "")]).strip(", ")
            _add_run(p, deg, size=11, font_name="Garamond")
            if ed.get("honors"):
                _add_run(p, f"; {ed['honors']}", size=11, font_name="Garamond")
            p.paragraph_format.space_after = Pt(4)

    if tailored.get("skills") or tailored.get("certifications"):
        _docx_section_heading(doc, "SKILLS & CERTIFICATIONS", font_name="Garamond",
                              size=12, centered=True, underline=True)
        skills = tailored.get("skills") or {}
        for label, items in [
            ("Core Competencies", skills.get("core_competencies", [])),
            ("Tools & Platforms", skills.get("tools_platforms", [])),
            ("Additional", skills.get("additional", [])),
        ]:
            if items:
                p = doc.add_paragraph()
                _add_run(p, f"{label}: ", bold=True, size=11, font_name="Garamond")
                _add_run(p, ", ".join(items), size=11, font_name="Garamond")
                p.paragraph_format.space_after = Pt(1)
        for c in (tailored.get("certifications") or []):
            p = doc.add_paragraph()
            bits = [c.get("name", "")]
            if c.get("issuer"):
                bits.append(c["issuer"])
            if c.get("year"):
                bits.append(str(c["year"]))
            _add_run(p, ", ".join(bits), size=11, font_name="Garamond")
            p.paragraph_format.space_after = Pt(1)

    doc.save(str(out_path))
    return out_path


def _inspect_user_docx(source_path: Path) -> tuple[str, int, int]:
    """Return (body_font, body_size, heading_size) inferred from a .docx."""
    body_font, body_size, heading_size = "Calibri", 11, 13
    if not source_path.exists() or source_path.suffix.lower() != ".docx":
        return body_font, body_size, heading_size
    try:
        src_doc = Document(str(source_path))
        sizes, fonts = [], []
        for para in src_doc.paragraphs[:30]:
            for run in para.runs[:5]:
                if run.font.size:
                    sizes.append(run.font.size.pt)
                if run.font.name:
                    fonts.append(run.font.name)
        if fonts:
            from collections import Counter
            for name, _ in Counter(fonts).most_common():
                if name in {"Calibri", "Arial", "Garamond",
                            "Georgia", "Times New Roman", "Helvetica"}:
                    body_font = name
                    break
        if sizes:
            body_sizes = [s for s in sizes if 9 <= s <= 13]
            if body_sizes:
                body_size = int(round(sum(body_sizes) / len(body_sizes)))
                heading_size = body_size + 2
    except Exception:
        pass
    return body_font, body_size, heading_size


def _build_mirror_user_docx(profile: dict, tailored: dict, out_path: Path,
                            source_docx_path: str | Path | None = None) -> Path:
    body_font, body_size, heading_size = "Calibri", 11, 13
    if source_docx_path:
        body_font, body_size, heading_size = _inspect_user_docx(Path(source_docx_path))

    doc = Document()
    _set_margins(doc, 0.6, 0.6, 0.75, 0.75)

    p = doc.add_paragraph()
    _add_run(p, tailored["name"], bold=True, size=max(body_size + 6, 16), font_name=body_font)

    p = doc.add_paragraph()
    _add_run(p, tailored.get("contact_line", ""), size=body_size, font_name=body_font)
    p.paragraph_format.space_after = Pt(10)

    if tailored.get("summary"):
        _docx_section_heading(doc, "SUMMARY", font_name=body_font, size=heading_size)
        p = doc.add_paragraph()
        _add_run(p, tailored["summary"], size=body_size, font_name=body_font)
        p.paragraph_format.space_after = Pt(8)

    if tailored.get("skills"):
        _docx_section_heading(doc, "SKILLS", font_name=body_font, size=heading_size)
        for label, items in [
            ("Core Competencies", tailored["skills"].get("core_competencies", [])),
            ("Tools & Platforms", tailored["skills"].get("tools_platforms", [])),
            ("Additional", tailored["skills"].get("additional", [])),
        ]:
            if not items:
                continue
            p = doc.add_paragraph()
            _add_run(p, f"{label}: ", bold=True, size=body_size, font_name=body_font)
            _add_run(p, ", ".join(items), size=body_size, font_name=body_font)
            p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if tailored.get("experience"):
        _docx_section_heading(doc, "EXPERIENCE", font_name=body_font, size=heading_size)
        for exp in tailored["experience"]:
            p = doc.add_paragraph()
            _add_run(p, exp.get("title", ""), bold=True, size=body_size, font_name=body_font)
            _add_run(p, f"  |  {exp.get('company', '')}", size=body_size, font_name=body_font)
            if exp.get("location"):
                _add_run(p, f"  |  {exp['location']}", size=body_size, font_name=body_font)
            dr = _date_range(exp)
            if dr:
                _add_run(p, f"  |  {dr}", size=body_size, font_name=body_font)
            p.paragraph_format.space_after = Pt(2)

            for bullet in exp.get("bullets", []) or []:
                bp = doc.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=body_size, font_name=body_font)
                bp.paragraph_format.space_after = Pt(1)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if tailored.get("education"):
        _docx_section_heading(doc, "EDUCATION", font_name=body_font, size=heading_size)
        for ed in tailored["education"]:
            p = doc.add_paragraph()
            parts = [ed.get("degree", "")]
            if ed.get("major"):
                parts.append(ed["major"])
            _add_run(p, ", ".join([x for x in parts if x]), bold=True,
                     size=body_size, font_name=body_font)
            _add_run(p, f"  |  {ed.get('school', '')}", size=body_size, font_name=body_font)
            if ed.get("year"):
                _add_run(p, f"  |  {ed['year']}", size=body_size, font_name=body_font)

    if tailored.get("certifications"):
        _docx_section_heading(doc, "CERTIFICATIONS", font_name=body_font, size=heading_size)
        for c in tailored["certifications"]:
            p = doc.add_paragraph()
            _add_run(p, c.get("name", ""), bold=True, size=body_size, font_name=body_font)
            tail = []
            if c.get("issuer"):
                tail.append(c["issuer"])
            if c.get("year"):
                tail.append(str(c["year"]))
            if tail:
                _add_run(p, "  |  " + "  |  ".join(tail), size=body_size, font_name=body_font)

    doc.save(str(out_path))
    return out_path


# ===========================================================================
# PDF BUILDERS (reportlab direct — no docx intermediate)
# ===========================================================================

FONTS_DIR = Path(__file__).resolve().parent.parent / "assets" / "fonts"
_FONTS_REGISTERED = False


def _register_fonts() -> None:
    """Register bundled OFL TTF fonts with ReportLab. Safe to call repeatedly."""
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    for face, filename in [
        ("EBGaramond", "EBGaramond-Regular.ttf"),
        ("EBGaramond-Bold", "EBGaramond-Bold.ttf"),
        ("Carlito", "Carlito-Regular.ttf"),
        ("Carlito-Bold", "Carlito-Bold.ttf"),
    ]:
        path = FONTS_DIR / filename
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont(face, str(path)))
            except Exception:
                pass
    for family, regular, bold in [
        ("EBGaramond", "EBGaramond", "EBGaramond-Bold"),
        ("Carlito", "Carlito", "Carlito-Bold"),
    ]:
        try:
            registerFontFamily(family, normal=regular, bold=bold,
                               italic=regular, boldItalic=bold)
        except Exception:
            pass
    _FONTS_REGISTERED = True


def _resolve_font(preferred: str) -> tuple[str, str]:
    """Map a requested font family to (regular_face, bold_face) that reportlab
    can actually render. Falls back to built-ins if bundled TTFs are missing."""
    from reportlab.pdfbase import pdfmetrics
    mapping = {
        "Garamond": ("EBGaramond", "EBGaramond-Bold"),
        "Calibri": ("Carlito", "Carlito-Bold"),
    }
    regular, bold = mapping.get(preferred, (preferred, preferred + "-Bold"))
    registered = set(pdfmetrics.getRegisteredFontNames())
    if regular in registered and bold in registered:
        return regular, bold
    # Serif fallback: use reportlab's built-in Times
    if preferred in ("Garamond", "Georgia", "Times New Roman"):
        return "Times-Roman", "Times-Bold"
    # Sans fallback: Helvetica (always built in)
    return "Helvetica", "Helvetica-Bold"


def _esc(s: str) -> str:
    """Escape text for reportlab Paragraph (XML-ish markup)."""
    if s is None:
        return ""
    return (str(s).replace("&", "&amp;")
                  .replace("<", "&lt;")
                  .replace(">", "&gt;"))


def _pdf_styles(regular: str, bold: str, body_size: int = 11,
                heading_size: int = 13, name_size: int = 18,
                centered_headings: bool = False, underline_headings: bool = False):
    """Build the full set of ParagraphStyles used by the PDF templates."""
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.lib.styles import ParagraphStyle

    leading = body_size + 3

    name = ParagraphStyle(
        "name", fontName=bold, fontSize=name_size,
        leading=name_size + 2, spaceAfter=2,
        alignment=TA_CENTER if centered_headings else TA_LEFT,
    )
    contact = ParagraphStyle(
        "contact", fontName=regular, fontSize=body_size,
        leading=leading, spaceAfter=10,
        alignment=TA_CENTER if centered_headings else TA_LEFT,
    )
    heading = ParagraphStyle(
        "heading", fontName=bold, fontSize=heading_size,
        leading=heading_size + 2, spaceBefore=8, spaceAfter=4,
        alignment=TA_CENTER if centered_headings else TA_LEFT,
    )
    body = ParagraphStyle(
        "body", fontName=regular, fontSize=body_size,
        leading=leading, spaceAfter=2,
    )
    body_bold = ParagraphStyle(
        "body_bold", parent=body, fontName=bold,
    )
    body_right = ParagraphStyle(
        "body_right", parent=body, alignment=TA_RIGHT,
    )
    body_bold_right = ParagraphStyle(
        "body_bold_right", parent=body_bold, alignment=TA_RIGHT,
    )
    # Hanging-indent bullet: hanging lets continuation lines align past the glyph
    bullet = ParagraphStyle(
        "bullet", parent=body,
        leftIndent=18, bulletIndent=6, spaceBefore=0, spaceAfter=1,
    )
    return {
        "name": name, "contact": contact, "heading": heading,
        "body": body, "body_bold": body_bold,
        "body_right": body_right, "body_bold_right": body_bold_right,
        "bullet": bullet,
        "_underline_headings": underline_headings,
    }


def _heading(text: str, styles: dict):
    """Render a section heading, underlining its text if the template wants it."""
    from reportlab.platypus import Paragraph
    inner = _esc(text)
    if styles.get("_underline_headings"):
        inner = f"<u>{inner}</u>"
    return Paragraph(inner, styles["heading"])


def _two_col(left_flow, right_flow, left_w, right_w):
    """Borderless single-row table: left content + right-aligned content on one line.
    Works inside a Flowable story — lets us place 'Company / Jan 2019 – Present'
    where the date is flush to the right margin. Extracts as one line for ATS."""
    from reportlab.platypus import Table, TableStyle
    t = Table([[left_flow, right_flow]], colWidths=[left_w, right_w])
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    return t


def _bullet_para(text: str, styles: dict):
    from reportlab.platypus import Paragraph
    return Paragraph(_esc(text), styles["bullet"], bulletText="•")


def _render_skills_section(story, tailored, styles, heading_text="SKILLS"):
    from reportlab.platypus import Paragraph, Spacer
    skills = tailored.get("skills") or {}
    if not any(skills.get(k) for k in ("core_competencies", "tools_platforms", "additional")):
        return
    story.append(_heading(heading_text, styles))
    for label, items in [
        ("Core Competencies", skills.get("core_competencies", [])),
        ("Tools & Platforms", skills.get("tools_platforms", [])),
        ("Additional", skills.get("additional", [])),
    ]:
        if not items:
            continue
        text = f"<b>{_esc(label)}:</b> {_esc(', '.join(items))}"
        story.append(Paragraph(text, styles["body"]))
    story.append(Spacer(1, 4))


def _render_certifications_block(story, tailored, styles):
    from reportlab.platypus import Paragraph
    for c in tailored.get("certifications") or []:
        bits = [_esc(c.get("name", ""))]
        tail = []
        if c.get("issuer"):
            tail.append(_esc(c["issuer"]))
        if c.get("year"):
            tail.append(_esc(c["year"]))
        line = f"<b>{bits[0]}</b>"
        if tail:
            line += "  |  " + "  |  ".join(tail)
        story.append(Paragraph(line, styles["body"]))


def _build_clean_modern_pdf(profile: dict, tailored: dict, out_path: Path) -> Path:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, KeepTogether

    _register_fonts()
    regular, bold = _resolve_font("Calibri")

    left_m = right_m = 0.7 * inch
    top_m = bottom_m = 0.6 * inch
    page_w = LETTER[0] - left_m - right_m

    doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=left_m, rightMargin=right_m,
        topMargin=top_m, bottomMargin=bottom_m,
        title=f"Resume — {tailored.get('name', '')}",
        author=tailored.get("name", ""),
    )
    styles = _pdf_styles(regular, bold, body_size=11, heading_size=13, name_size=18)

    story = []
    story.append(Paragraph(_esc(tailored["name"]), styles["name"]))
    story.append(Paragraph(_esc(tailored.get("contact_line", "")), styles["contact"]))

    if tailored.get("summary"):
        story.append(_heading("SUMMARY", styles))
        story.append(Paragraph(_esc(tailored["summary"]), styles["body"]))
        story.append(Spacer(1, 6))

    _render_skills_section(story, tailored, styles)

    if tailored.get("experience"):
        story.append(_heading("EXPERIENCE", styles))
        for exp in tailored["experience"]:
            header_parts = [f"<b>{_esc(exp.get('title', ''))}</b>"]
            if exp.get("company"):
                header_parts.append(_esc(exp["company"]))
            if exp.get("location"):
                header_parts.append(_esc(exp["location"]))
            dr = _date_range(exp)
            if dr:
                header_parts.append(_esc(dr))
            header = Paragraph("  |  ".join(header_parts), styles["body"])
            bullets = [_bullet_para(b, styles) for b in exp.get("bullets") or []]
            story.append(KeepTogether([header, *bullets]))
            story.append(Spacer(1, 4))

    if tailored.get("education"):
        story.append(_heading("EDUCATION", styles))
        for ed in tailored["education"]:
            parts = [ed.get("degree", "")]
            if ed.get("major"):
                parts.append(ed["major"])
            left = _esc(", ".join([x for x in parts if x]))
            line = f"<b>{left}</b>  |  {_esc(ed.get('school', ''))}"
            if ed.get("year"):
                line += f"  |  {_esc(ed['year'])}"
            if ed.get("honors"):
                line += f"  |  {_esc(ed['honors'])}"
            story.append(Paragraph(line, styles["body"]))

    if tailored.get("certifications"):
        story.append(_heading("CERTIFICATIONS", styles))
        _render_certifications_block(story, tailored, styles)

    doc.build(story)
    return out_path


def _build_harvard_pdf(profile: dict, tailored: dict, out_path: Path) -> Path:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, KeepTogether

    _register_fonts()
    regular, bold = _resolve_font("Garamond")

    left_m = right_m = 0.9 * inch
    top_m = bottom_m = 0.75 * inch
    page_w = LETTER[0] - left_m - right_m
    # Left column takes ~70%, right ~30% — enough for "Jan 2019 – Present"
    left_w = page_w * 0.68
    right_w = page_w - left_w

    doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=left_m, rightMargin=right_m,
        topMargin=top_m, bottomMargin=bottom_m,
        title=f"Resume — {tailored.get('name', '')}",
        author=tailored.get("name", ""),
    )
    styles = _pdf_styles(
        regular, bold, body_size=11, heading_size=12, name_size=16,
        centered_headings=True, underline_headings=True,
    )

    story = []
    story.append(Paragraph(_esc(tailored["name"]), styles["name"]))
    story.append(Paragraph(_esc(tailored.get("contact_line", "")), styles["contact"]))

    if tailored.get("summary"):
        story.append(_heading("SUMMARY", styles))
        story.append(Paragraph(_esc(tailored["summary"]), styles["body"]))
        story.append(Spacer(1, 4))

    if tailored.get("experience"):
        story.append(_heading("EXPERIENCE", styles))
        for exp in tailored["experience"]:
            # Row 1: company (bold, left) | dates (right-aligned)
            company = Paragraph(f"<b>{_esc(exp.get('company', ''))}</b>", styles["body_bold"])
            dr = _date_range(exp)
            dates = Paragraph(_esc(dr), styles["body_right"]) if dr else Paragraph("", styles["body_right"])
            row1 = _two_col(company, dates, left_w, right_w)

            # Row 2: title (italic-ish, here just regular) | location (right-aligned)
            title = Paragraph(_esc(exp.get("title", "")), styles["body"])
            loc = Paragraph(_esc(exp.get("location", "")), styles["body_right"])
            row2 = _two_col(title, loc, left_w, right_w)

            bullets = [_bullet_para(b, styles) for b in exp.get("bullets") or []]
            story.append(KeepTogether([row1, row2, Spacer(1, 1), *bullets]))
            story.append(Spacer(1, 4))

    if tailored.get("education"):
        story.append(_heading("EDUCATION", styles))
        for ed in tailored["education"]:
            school = Paragraph(f"<b>{_esc(ed.get('school', ''))}</b>", styles["body_bold"])
            year = Paragraph(_esc(ed.get("year", "") or ""), styles["body_right"])
            story.append(_two_col(school, year, left_w, right_w))
            deg = ", ".join([ed.get("degree", ""), ed.get("major", "")]).strip(", ")
            line = _esc(deg)
            if ed.get("honors"):
                line += f"; {_esc(ed['honors'])}"
            story.append(Paragraph(line, styles["body"]))
            story.append(Spacer(1, 4))

    if tailored.get("skills") or tailored.get("certifications"):
        story.append(_heading("SKILLS & CERTIFICATIONS", styles))
        skills = tailored.get("skills") or {}
        for label, items in [
            ("Core Competencies", skills.get("core_competencies", [])),
            ("Tools & Platforms", skills.get("tools_platforms", [])),
            ("Additional", skills.get("additional", [])),
        ]:
            if items:
                story.append(Paragraph(
                    f"<b>{_esc(label)}:</b> {_esc(', '.join(items))}",
                    styles["body"],
                ))
        for c in tailored.get("certifications") or []:
            bits = [c.get("name", "")]
            if c.get("issuer"):
                bits.append(c["issuer"])
            if c.get("year"):
                bits.append(str(c["year"]))
            story.append(Paragraph(_esc(", ".join(bits)), styles["body"]))

    doc.build(story)
    return out_path


def _build_mirror_user_pdf(profile: dict, tailored: dict, out_path: Path,
                           source_docx_path: str | Path | None = None) -> Path:
    # Inspect the user's uploaded docx (if any) to pick font + size, then
    # render in the Clean Modern PDF layout with their choices applied.
    body_font_pref, body_size, heading_size = "Calibri", 11, 13
    if source_docx_path:
        body_font_pref, body_size, heading_size = _inspect_user_docx(Path(source_docx_path))

    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, KeepTogether

    _register_fonts()
    regular, bold = _resolve_font(body_font_pref)

    left_m = right_m = 0.75 * inch
    top_m = bottom_m = 0.6 * inch

    doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=left_m, rightMargin=right_m,
        topMargin=top_m, bottomMargin=bottom_m,
        title=f"Resume — {tailored.get('name', '')}",
        author=tailored.get("name", ""),
    )
    styles = _pdf_styles(
        regular, bold, body_size=body_size, heading_size=heading_size,
        name_size=max(body_size + 6, 16),
    )

    story = []
    story.append(Paragraph(_esc(tailored["name"]), styles["name"]))
    story.append(Paragraph(_esc(tailored.get("contact_line", "")), styles["contact"]))

    if tailored.get("summary"):
        story.append(_heading("SUMMARY", styles))
        story.append(Paragraph(_esc(tailored["summary"]), styles["body"]))
        story.append(Spacer(1, 6))

    _render_skills_section(story, tailored, styles)

    if tailored.get("experience"):
        story.append(_heading("EXPERIENCE", styles))
        for exp in tailored["experience"]:
            header_parts = [f"<b>{_esc(exp.get('title', ''))}</b>"]
            if exp.get("company"):
                header_parts.append(_esc(exp["company"]))
            if exp.get("location"):
                header_parts.append(_esc(exp["location"]))
            dr = _date_range(exp)
            if dr:
                header_parts.append(_esc(dr))
            header = Paragraph("  |  ".join(header_parts), styles["body"])
            bullets = [_bullet_para(b, styles) for b in exp.get("bullets") or []]
            story.append(KeepTogether([header, *bullets]))
            story.append(Spacer(1, 4))

    if tailored.get("education"):
        story.append(_heading("EDUCATION", styles))
        for ed in tailored["education"]:
            parts = [ed.get("degree", "")]
            if ed.get("major"):
                parts.append(ed["major"])
            left = _esc(", ".join([x for x in parts if x]))
            line = f"<b>{left}</b>  |  {_esc(ed.get('school', ''))}"
            if ed.get("year"):
                line += f"  |  {_esc(ed['year'])}"
            story.append(Paragraph(line, styles["body"]))

    if tailored.get("certifications"):
        story.append(_heading("CERTIFICATIONS", styles))
        _render_certifications_block(story, tailored, styles)

    doc.build(story)
    return out_path


# ===========================================================================
# Public API — dispatch by output file extension
# ===========================================================================

def _dispatch(docx_builder, pdf_builder, profile, tailored, out_path, **kwargs):
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = out_path.suffix.lower()
    if suffix == ".pdf":
        return pdf_builder(profile, tailored, out_path, **kwargs)
    # Default to docx for .docx or unknown extension
    return docx_builder(profile, tailored, out_path, **kwargs)


def build_clean_modern(profile: dict, tailored: dict, out_path: str | Path) -> Path:
    """Render the Clean Modern template. Output is .docx or .pdf based on suffix."""
    return _dispatch(_build_clean_modern_docx, _build_clean_modern_pdf,
                     profile, tailored, out_path)


def build_harvard(profile: dict, tailored: dict, out_path: str | Path) -> Path:
    """Render the Harvard template. Output is .docx or .pdf based on suffix."""
    return _dispatch(_build_harvard_docx, _build_harvard_pdf,
                     profile, tailored, out_path)


def build_mirror_user(profile: dict, tailored: dict, out_path: str | Path,
                      source_docx_path: str | Path | None = None) -> Path:
    """Render the Mirror-User template. Output is .docx or .pdf based on suffix."""
    return _dispatch(_build_mirror_user_docx, _build_mirror_user_pdf,
                     profile, tailored, out_path,
                     source_docx_path=source_docx_path)


def build_resume(template: str, profile: dict, tailored: dict,
                 out_path: str | Path,
                 source_docx_path: str | Path | None = None) -> Path:
    """Render `tailored` via the named template to out_path. Output format is
    determined by the file extension (.docx → python-docx, .pdf → reportlab)."""
    if template == "clean_modern":
        return build_clean_modern(profile, tailored, out_path)
    if template == "harvard":
        return build_harvard(profile, tailored, out_path)
    if template == "mirror_user":
        return build_mirror_user(profile, tailored, out_path, source_docx_path)
    return build_clean_modern(profile, tailored, out_path)


# ===========================================================================
# Fabrication check (unchanged)
# ===========================================================================

_NUMBER_RE = re.compile(r"\b\d+(?:\.\d+)?(?:%|\b)|\$[\d,]+[KMBkmb]?")


def fabrication_check(tailored: dict, profile: dict) -> list[dict]:
    """Flag claims in the tailored resume that can't be traced to the profile.

    Returns a list of findings. Each finding is:
        {"location": str, "claim": str, "issue": str}

    Findings are advisory — Stage 6 should show them to the user and ask
    for verification before saving the resume.
    """
    from scripts.profile import has_evidence

    findings = []

    profile_text_chunks = []
    for exp in profile.get("experience", []):
        for b in exp.get("bullets", []) or []:
            profile_text_chunks.append(b)
        for t in exp.get("technologies", []) or []:
            profile_text_chunks.append(t)
        profile_text_chunks.append(exp.get("title", ""))
        profile_text_chunks.append(exp.get("company", ""))
    for p in profile.get("projects", []) or []:
        profile_text_chunks.append(p.get("description", ""))
        for o in p.get("outcomes", []) or []:
            profile_text_chunks.append(o)
        for t in p.get("technologies", []) or []:
            profile_text_chunks.append(t)
    profile_text = " | ".join(profile_text_chunks).lower()

    profile_numbers = set()
    for chunk in profile_text_chunks:
        profile_numbers.update(m.group(0).lower() for m in _NUMBER_RE.finditer(chunk or ""))
    yrs = profile.get("years_experience_total")
    if yrs is not None:
        profile_numbers.add(str(int(yrs)) if yrs == int(yrs) else str(yrs))
        profile_numbers.add(str(yrs))
        for delta in (-1, 0, 1):
            v = int(yrs) + delta
            if v >= 0:
                profile_numbers.add(str(v))
    for ed in profile.get("education", []) or []:
        if ed.get("year"):
            profile_numbers.add(str(ed["year"]).lower())
    for exp in profile.get("experience", []) or []:
        for k in ("start", "end"):
            v = exp.get(k)
            if v:
                m = re.match(r"(\d{4})", str(v))
                if m:
                    profile_numbers.add(m.group(1))

    summary = tailored.get("summary", "") or ""
    for m in _NUMBER_RE.finditer(summary):
        n = m.group(0).lower()
        if n not in profile_numbers and not _is_trivially_safe_number(n):
            findings.append({
                "location": "summary",
                "claim": n,
                "issue": "number not found in profile — is this based on a real figure?",
            })
    m = re.search(r"team of (\d+)", summary.lower())
    if m:
        team_claim = f"team of {m.group(1)}"
        if team_claim not in profile_text and m.group(1) not in profile_numbers:
            findings.append({
                "location": "summary",
                "claim": team_claim,
                "issue": "team size not found in profile",
            })

    for exp_idx, exp in enumerate(tailored.get("experience", []) or []):
        company = exp.get("company", "")
        for bi, bullet in enumerate(exp.get("bullets", []) or []):
            for m in _NUMBER_RE.finditer(bullet):
                n = m.group(0).lower()
                if n not in profile_numbers and not _is_trivially_safe_number(n):
                    findings.append({
                        "location": f"experience #{exp_idx+1} ({company}), bullet #{bi+1}",
                        "claim": n,
                        "issue": "number not in profile — invented?",
                    })
            m = re.search(r"team of (\d+)", bullet.lower())
            if m:
                team_claim = f"team of {m.group(1)}"
                if team_claim not in profile_text and m.group(1) not in profile_numbers:
                    findings.append({
                        "location": f"experience #{exp_idx+1} ({company}), bullet #{bi+1}",
                        "claim": team_claim,
                        "issue": "team size not found in profile",
                    })

    skills_obj = tailored.get("skills") or {}
    for bucket_name in ("core_competencies", "tools_platforms", "additional"):
        for s in skills_obj.get(bucket_name, []) or []:
            if not has_evidence(profile, s):
                findings.append({
                    "location": f"skills.{bucket_name}",
                    "claim": s,
                    "issue": "no evidence for this skill in profile — remove or add evidence",
                })

    tailored_edu = {(e.get("school", "").lower(), (e.get("degree", "") or "").lower())
                    for e in tailored.get("education", []) or []}
    profile_edu = {(e.get("school", "").lower(), (e.get("degree", "") or "").lower())
                   for e in profile.get("education", []) or []}
    for entry in tailored_edu:
        if entry not in profile_edu and entry[0]:
            findings.append({
                "location": "education",
                "claim": f"{entry[1]} from {entry[0]}",
                "issue": "education entry not found in canonical profile",
            })

    tailored_certs = {(c.get("name", "").lower()) for c in tailored.get("certifications", []) or []}
    profile_certs = {(c.get("name", "").lower()) for c in profile.get("certifications", []) or []}
    for name in tailored_certs:
        if name and name not in profile_certs:
            findings.append({
                "location": "certifications",
                "claim": name,
                "issue": "certification not in canonical profile",
            })

    return findings


def _is_trivially_safe_number(s: str) -> bool:
    s = s.strip().lower().replace(",", "").replace("%", "").replace("$", "")
    s = s.rstrip("kmb")
    if not s:
        return True
    try:
        n = float(s)
    except ValueError:
        return False
    if 1990 <= n <= 2099:
        return True
    return False

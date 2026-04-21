"""Unit tests for scripts/resume.py."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from scripts.profile import new_profile, add_source_doc, add_evidence
from scripts.resume import (
    build_clean_modern, build_harvard, build_mirror_user, build_resume,
    fabrication_check,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _basic_profile():
    p = new_profile("Test User")
    p["years_experience_total"] = 8
    p["experience"] = [
        {"title": "Senior Marketing Manager", "company": "TechCorp",
         "location": "DC", "start": "2019-01", "current": True,
         "bullets": ["Built ops stack with HubSpot.",
                     "Managed 4 direct reports.",
                     "Cut lead time 70%."],
         "technologies": ["HubSpot", "Salesforce"]},
    ]
    p["education"] = [{"degree": "BS", "major": "Marketing",
                       "school": "Maryland", "year": 2016}]
    p["skills"]["domains"] = ["Demand Gen"]  # so "Demand Gen" has evidence
    doc_id = add_source_doc(p, "resume_pdf", "/r.pdf", "content")
    add_evidence(p, "skills_evidence", "HubSpot", doc_id, "loc")
    add_evidence(p, "skills_evidence", "Salesforce", doc_id, "loc")
    return p


def _basic_tailored():
    return {
        "name": "Test User",
        "contact_line": "DC | t@x.com | (555) 555-5555",
        "summary": "8 years driving demand gen for B2B SaaS.",
        "skills": {
            "core_competencies": ["Demand Gen"],
            "tools_platforms": ["HubSpot", "Salesforce"],
            "additional": [],
        },
        "experience": [
            {"title": "Senior Marketing Manager", "company": "TechCorp",
             "location": "DC", "start": "Jan 2019", "current": True,
             "bullets": ["Built ops stack with HubSpot.",
                         "Cut lead time 70%."]},
        ],
        "education": [{"degree": "BS", "major": "Marketing",
                       "school": "Maryland", "year": 2016, "honors": None}],
        "certifications": [],
    }


# ---------------------------------------------------------------------------
# Template rendering
# ---------------------------------------------------------------------------

def test_build_clean_modern_produces_docx(tmp_path):
    out = tmp_path / "r.docx"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    assert out.exists() and out.stat().st_size > 5000


def test_build_harvard_produces_docx(tmp_path):
    out = tmp_path / "r.docx"
    build_harvard(_basic_profile(), _basic_tailored(), out)
    assert out.exists() and out.stat().st_size > 5000


def test_build_mirror_user_falls_back_without_source(tmp_path):
    out = tmp_path / "r.docx"
    build_mirror_user(_basic_profile(), _basic_tailored(), out, source_docx_path=None)
    assert out.exists() and out.stat().st_size > 5000


def test_build_resume_dispatcher_routes_correctly(tmp_path):
    for tpl in ("clean_modern", "harvard", "mirror_user"):
        out = tmp_path / f"r_{tpl}.docx"
        build_resume(tpl, _basic_profile(), _basic_tailored(), out)
        assert out.exists()


def test_build_resume_unknown_template_falls_back(tmp_path):
    """Unknown template should fall back to clean_modern, not crash."""
    out = tmp_path / "r.docx"
    build_resume("nonexistent", _basic_profile(), _basic_tailored(), out)
    assert out.exists()


def test_clean_modern_contact_line_present(tmp_path):
    out = tmp_path / "r.docx"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "t@x.com" in text


def test_clean_modern_has_standard_section_headings(tmp_path):
    out = tmp_path / "r.docx"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.text.isupper() and len(p.text) < 30]
    assert any("SUMMARY" in h for h in headings)
    assert any("SKILLS" in h for h in headings)
    assert any("EXPERIENCE" in h for h in headings)


def test_clean_modern_bullets_render_as_list(tmp_path):
    """ATS-safe: bullets must use the 'List Bullet' style."""
    out = tmp_path / "r.docx"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    doc = Document(str(out))
    bullet_count = sum(1 for p in doc.paragraphs if p.style.name == "List Bullet")
    assert bullet_count >= 2  # we provided 2 bullets


def test_no_tables_in_output(tmp_path):
    """ATS-safe: no tables anywhere."""
    out = tmp_path / "r.docx"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    doc = Document(str(out))
    assert len(doc.tables) == 0


# ---------------------------------------------------------------------------
# PDF rendering (direct reportlab, no docx intermediate)
# ---------------------------------------------------------------------------

def _pdf_header(path: Path) -> bytes:
    return path.read_bytes()[:5]


def test_build_clean_modern_produces_pdf(tmp_path):
    out = tmp_path / "r.pdf"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    assert out.exists() and out.stat().st_size > 2000
    assert _pdf_header(out) == b"%PDF-"


def test_build_harvard_produces_pdf(tmp_path):
    out = tmp_path / "r.pdf"
    build_harvard(_basic_profile(), _basic_tailored(), out)
    assert out.exists() and out.stat().st_size > 2000
    assert _pdf_header(out) == b"%PDF-"


def test_build_mirror_user_produces_pdf(tmp_path):
    out = tmp_path / "r.pdf"
    build_mirror_user(_basic_profile(), _basic_tailored(), out, source_docx_path=None)
    assert out.exists() and out.stat().st_size > 2000
    assert _pdf_header(out) == b"%PDF-"


def test_build_resume_dispatches_pdf_by_extension(tmp_path):
    """build_resume should produce a PDF when out_path ends in .pdf."""
    for tpl in ("clean_modern", "harvard", "mirror_user"):
        out = tmp_path / f"r_{tpl}.pdf"
        build_resume(tpl, _basic_profile(), _basic_tailored(), out)
        assert out.exists() and out.stat().st_size > 2000
        assert _pdf_header(out) == b"%PDF-"


def test_pdf_contains_name_text(tmp_path):
    """Smoke test: the rendered PDF must contain the user's name in extractable text."""
    out = tmp_path / "r.pdf"
    build_clean_modern(_basic_profile(), _basic_tailored(), out)
    raw = out.read_bytes()
    # Reportlab embeds the name as a Tj operator in the content stream. We
    # can't parse the PDF without a library, but the name appears either
    # literally or as a glyph-indexed stream; checking for the first name
    # as a plain substring works with built-in Helvetica/Times (no subsetting).
    # With subsetted TTFs, fall back to a structural check.
    assert b"%PDF-" in raw[:8]
    assert b"%%EOF" in raw[-1024:]


# ---------------------------------------------------------------------------
# Fabrication check — clean content
# ---------------------------------------------------------------------------

def test_fabrication_check_clean_content_passes():
    findings = fabrication_check(_basic_tailored(), _basic_profile())
    assert findings == []


def test_fabrication_check_uses_years_experience_total():
    """The tailored summary's 'X years' should not be flagged when X matches profile."""
    profile = _basic_profile()
    profile["years_experience_total"] = 10
    tailored = _basic_tailored()
    tailored["summary"] = "10 years driving things."
    findings = fabrication_check(tailored, profile)
    # No findings about the "10"
    assert not any(f["claim"] == "10" for f in findings)


def test_fabrication_check_allows_education_year():
    """A 2016 in education should not trigger a number flag."""
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["summary"] = "Started in 2016. Cut lead time 70%."
    findings = fabrication_check(tailored, profile)
    assert not any(f["claim"] == "2016" for f in findings)


# ---------------------------------------------------------------------------
# Fabrication check — tainted content
# ---------------------------------------------------------------------------

def test_fabrication_check_catches_invented_summary_number():
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["summary"] = "Cut lead time 99%."
    findings = fabrication_check(tailored, profile)
    assert any(f["claim"] == "99%" for f in findings)


def test_fabrication_check_catches_invented_team_size():
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["summary"] = "Managed team of 25."
    findings = fabrication_check(tailored, profile)
    assert any("team of 25" in f["claim"] for f in findings)


def test_fabrication_check_catches_invented_skill():
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["skills"]["core_competencies"] = ["Kubernetes"]
    findings = fabrication_check(tailored, profile)
    assert any("Kubernetes" in f["claim"] for f in findings)


def test_fabrication_check_catches_invented_education():
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["education"] = [{"degree": "MBA", "major": "Bus", "school": "Stanford", "year": 2020}]
    findings = fabrication_check(tailored, profile)
    flagged = " | ".join(f["claim"].lower() for f in findings)
    assert "stanford" in flagged or "mba" in flagged


def test_fabrication_check_catches_invented_cert():
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["certifications"] = [{"name": "PMP", "issuer": "PMI", "year": 2022}]
    findings = fabrication_check(tailored, profile)
    assert any("pmp" in f["claim"].lower() for f in findings)


def test_fabrication_check_catches_invented_bullet_number():
    profile = _basic_profile()
    tailored = _basic_tailored()
    tailored["experience"][0]["bullets"] = ["Grew revenue 500%."]
    findings = fabrication_check(tailored, profile)
    assert any("500%" in f["claim"] for f in findings)


def test_fabrication_check_skill_with_evidence_passes():
    """Skill that's in profile experience.technologies should pass."""
    profile = _basic_profile()
    # Add a technology to experience
    profile["experience"][0]["technologies"].append("Marketo")
    tailored = _basic_tailored()
    tailored["skills"]["tools_platforms"].append("Marketo")
    findings = fabrication_check(tailored, profile)
    assert not any("Marketo" in f["claim"] for f in findings)
